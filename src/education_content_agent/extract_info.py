import os
import fitz  # PyMuPDF para PDFs con texto
import easyocr  # OCR multilenguaje
from pdf2image import convert_from_path  # Convierte PDF en imágenes para OCR
from PIL import Image  # Manejo de imágenes
from docx import Document  # Para leer .docx
import numpy as np  # Para trabajar con arrays numpy

class DocumentExtractor:
    def __init__(self, idioma='es'):
        """
        Inicializa el extractor de documentos.
        
        Args:
            idioma (str): Código del idioma para OCR (default: 'es' para español)
        """
        self.reader = easyocr.Reader([idioma])

    def extraer_texto_docx(self, ruta_docx):
        """
        Extrae texto de un documento DOCX, incluyendo contenido de tablas.
        """
        doc = Document(ruta_docx)
        texto_completo = ""

        # Extraer texto de los párrafos normales
        for para in doc.paragraphs:
            texto_completo += para.text + "\n"

        # Extraer texto de las tablas
        for table in doc.tables:
            for fila in table.rows:
                for celda in fila.cells:
                    texto_completo += celda.text + " | "  # Separador entre columnas
                texto_completo += "\n"  # Salto de línea al final de cada fila
            texto_completo += "\n"  # Espacio extra entre tablas

        return texto_completo

    def extraer_texto_pdf(self, ruta_pdf):
        """
        Extrae texto de un PDF con PyMuPDF (si no está escaneado).
        Si no detecta texto, usa OCR.
        """
        texto_extraido = ""
        doc = fitz.open(ruta_pdf)

        for num_pagina in range(len(doc)):
            texto = doc[num_pagina].get_text("text")

            if texto.strip():  # Si hay texto, lo añadimos
                texto_extraido += texto + "\n\n"
            else:  # Si la página no tiene texto, usamos OCR
                print(f"⚠ Página {num_pagina + 1} parece ser escaneada, aplicando OCR...")
                imagenes = convert_from_path(ruta_pdf, first_page=num_pagina+1, last_page=num_pagina+1)
                
                for img in imagenes:
                    # Convertimos la imagen a array numpy si es necesario
                    img_array = np.array(img)
                    # Realizamos OCR con EasyOCR
                    resultados = self.reader.readtext(img_array)
                    # Extraemos el texto de los resultados
                    texto_ocr = '\n'.join([resultado[1] for resultado in resultados])
                    texto_extraido += texto_ocr + "\n\n"

        return texto_extraido

    def extraer_texto_txt(self, ruta_txt):
        """
        Extrae texto de archivos TXT.
        """
        with open(ruta_txt, "r", encoding="utf-8") as f:
            return f.read() + "\n\n"

    def procesar_carpeta(self, ruta_carpeta, tipos_archivo=None):
        """
        Extrae y concatena el texto de todos los archivos soportados en una carpeta.
        
        Args:
            ruta_carpeta (str): Ruta a la carpeta con los documentos
            tipos_archivo (list): Lista de extensiones a procesar (ej: ['.pdf', '.docx'])
                                Si es None, procesa todos los tipos soportados
        
        Returns:
            str: Texto extraído de todos los archivos concatenado
        """
        if tipos_archivo is None:
            tipos_archivo = ['.txt', '.docx', '.pdf']
        
        texto_completo = ""

        for archivo in os.listdir(ruta_carpeta):
            nombre, extension = os.path.splitext(archivo)
            if extension.lower() not in tipos_archivo:
                continue

            ruta_archivo = os.path.join(ruta_carpeta, archivo)
            
            if not os.path.isfile(ruta_archivo):
                continue

            try:
                if extension.lower() == '.txt':
                    print(f"📄 Procesando TXT: {archivo}")
                    texto = self.extraer_texto_txt(ruta_archivo)
                
                elif extension.lower() == '.docx':
                    print(f"📄 Procesando DOCX: {archivo}")
                    texto = self.extraer_texto_docx(ruta_archivo)
                
                elif extension.lower() == '.pdf':
                    print(f"📄 Procesando PDF: {archivo}")
                    texto = self.extraer_texto_pdf(ruta_archivo)
                
                texto_completo += f"\n\n--- Contenido de {archivo} ---\n\n"
                texto_completo += texto

            except Exception as e:
                print(f"❌ Error procesando {archivo}: {e}")
                texto_completo += f"\n\nERROR en {archivo}: {str(e)}\n\n"

        return texto_completo

# # Ruta de la carpeta donde están los archivos
# ruta_carpeta = "Docs"  # Cambia esta ruta según la ubicación de tus archivos

# # Extraer y mostrar el texto de todos los archivos en la carpeta
# extractor = DocumentExtractor()
# resultados = extractor.procesar_carpeta(ruta_carpeta)
# for archivo, texto in resultados.items():
#     print(f"Contenido de {archivo}:\n{texto}\n")